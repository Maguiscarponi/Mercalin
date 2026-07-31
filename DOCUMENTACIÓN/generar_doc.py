"""
Genera Analisis_Estrategico_PuntoSimple.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── helpers ────────────────────────────────────────────────────────────────

def rgb(r, g, b):
    return RGBColor(r, g, b)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            for key, val in kwargs[edge].items():
                tag.set(qn(f'w:{key}'), val)
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_run(para, text, bold=False, italic=False, size=11,
            color=None, font=None):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if font:
        run.font.name = font
    return run

def heading(doc, text, level=1, color=None, size=None, space_before=18, space_after=8):
    sizes = {1: 20, 2: 15, 3: 13, 4: 11}
    colors_default = {
        1: rgb(15, 23, 42),
        2: rgb(30, 58, 95),
        3: rgb(51, 65, 85),
        4: rgb(71, 85, 105),
    }
    para = doc.add_paragraph()
    para.style = f'Heading {level}'
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size or sizes[level])
    run.font.color.rgb = color or colors_default[level]
    run.font.name = 'Calibri'
    return para

def body(doc, text, space_before=2, space_after=6, color=None, size=10.5, italic=False):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.italic = italic
    if color:
        run.font.color.rgb = color
    else:
        run.font.color.rgb = rgb(71, 85, 105)
    return para

def bullet(doc, text, level=0, bold_prefix=None):
    para = doc.add_paragraph(style='List Bullet')
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    if bold_prefix:
        r1 = para.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.name = 'Calibri'
        r1.font.color.rgb = rgb(30, 58, 95)
        r2 = para.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.name = 'Calibri'
        r2.font.color.rgb = rgb(71, 85, 105)
    else:
        run = para.add_run(text)
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        run.font.color.rgb = rgb(71, 85, 105)
    return para

def colored_box(doc, title, text, bg='EFF6FF', title_color=None, text_color=None, icon=''):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg)
    cell.width = Inches(6.3)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after = Pt(2)
    p1.paragraph_format.left_indent = Cm(0.3)
    if icon:
        r0 = p1.add_run(icon + '  ')
        r0.font.size = Pt(12)
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = 'Calibri'
    r1.font.color.rgb = title_color or rgb(30, 64, 175)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.left_indent = Cm(0.3)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = text_color or rgb(30, 27, 75)
    doc.add_paragraph()  # spacing
    return tbl

def page_break(doc):
    doc.add_page_break()

def divider(doc):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para

def add_styled_table(doc, headers, rows, col_widths=None):
    """headers: list of str. rows: list of list of str/tuple (text, bold, bg)."""
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style = 'Table Grid'

    # header row
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, '0F172A')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.15)
        run = p.add_run(h.upper())
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.name = 'Calibri'
        run.font.color.rgb = rgb(255, 255, 255)

    # data rows
    for ri, row in enumerate(rows):
        tr = tbl.add_row()
        bg = 'F8FAFC' if ri % 2 == 1 else 'FFFFFF'
        for ci, cell_data in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.15)
            if isinstance(cell_data, tuple):
                text, bold, tag_bg, tag_fg = cell_data[0], cell_data[1], cell_data[2] if len(cell_data) > 2 else None, cell_data[3] if len(cell_data) > 3 else None
            else:
                text, bold, tag_bg, tag_fg = str(cell_data), (ci == 0), None, None

            if tag_bg:
                # draw as colored tag
                run = p.add_run(f' {text} ')
                run.font.size = Pt(9)
                run.font.name = 'Calibri'
                run.bold = True
                run.font.color.rgb = rgb(*tag_fg) if tag_fg else rgb(30, 64, 175)
            else:
                run = p.add_run(text)
                run.bold = bold
                run.font.size = Pt(9.5)
                run.font.name = 'Calibri'
                run.font.color.rgb = rgb(51, 65, 85)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in tbl.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return tbl

# ─── DOCUMENT ───────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

# ═══════════════════════════════════════ PORTADA ════════════════════════════

# Spacer
for _ in range(5):
    doc.add_paragraph()

# Empresa
p_company = doc.add_paragraph()
p_company.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_company.add_run('PUNTO SIMPLE POS')
run.bold = True
run.font.size = Pt(32)
run.font.name = 'Calibri'
run.font.color.rgb = rgb(15, 23, 42)

# Título
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_title.add_run('Análisis Estratégico Completo')
run.font.size = Pt(20)
run.font.name = 'Calibri'
run.font.color.rgb = rgb(99, 102, 241)

p_subtitle = doc.add_paragraph()
p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_subtitle.add_run('Diagnóstico del sistema · Oportunidades de mejora · Hoja de ruta')
run.font.size = Pt(13)
run.font.name = 'Calibri'
run.italic = True
run.font.color.rgb = rgb(100, 116, 139)

for _ in range(3):
    doc.add_paragraph()

# Meta-info table
meta_tbl = doc.add_table(rows=1, cols=4)
meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (label, value) in enumerate([
    ('VERSIÓN', '2.0 — 2026'),
    ('FECHA', 'Mayo 2026'),
    ('ESTADO', 'Diagnóstico'),
    ('CLASIFICACIÓN', 'Interno'),
]):
    cell = meta_tbl.cell(0, i)
    set_cell_bg(cell, '0F172A')
    cell.width = Cm(4)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    r1 = p.add_run(label + '\n')
    r1.font.size = Pt(7)
    r1.font.name = 'Calibri'
    r1.font.color.rgb = rgb(100, 116, 139)
    r1.bold = True
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = rgb(226, 232, 240)
    r2.bold = True
    p.paragraph_format.space_after = Pt(6)

doc.add_paragraph()

# Clasificación badge
p_badge = doc.add_paragraph()
p_badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_badge.add_run('  DOCUMENTO TÉCNICO CONFIDENCIAL  ')
run.bold = True
run.font.size = Pt(9)
run.font.name = 'Calibri'
run.font.color.rgb = rgb(165, 180, 252)

page_break(doc)

# ═══════════════════════════════════════ RESUMEN EJECUTIVO ══════════════════

heading(doc, 'Resumen Ejecutivo', level=1)
divider(doc)

colored_box(
    doc,
    '¿Dónde está parado el sistema hoy?',
    'Punto Simple POS es hoy una caja registradora digital bien ejecutada: '
    'vende, cobra, controla stock y maneja cuentas corrientes. Sin embargo, el sistema '
    'no le habla al dueño. No avisa cuando algo va mal. No detecta oportunidades. No aprende. '
    'Todo el valor de los datos acumulados (ventas históricas, patrones de comportamiento, '
    'márgenes reales) está dormido en la base de datos sin ser aprovechado.\n\n'
    'Este documento define la transformación de una herramienta de registro en un socio '
    'inteligente del negocio: que analiza, aprende, anticipa y recomienda acciones concretas.',
    bg='EFF6FF', title_color=rgb(30, 64, 175), text_color=rgb(30, 27, 75), icon='⚡'
)

body(doc, 'Métricas del análisis:', bold=False) if False else None

# Stats summary as mini table
stats_tbl = doc.add_table(rows=1, cols=4)
stats_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
for i, (num, lbl) in enumerate([
    ('15', 'Módulos funcionales\noperativos'),
    ('42', 'Features nuevas\nidentificadas'),
    ('23', 'Insights automáticos\nposibles'),
    ('4', 'Fases de\nimplementación'),
]):
    cell = stats_tbl.cell(0, i)
    set_cell_bg(cell, 'F8FAFC')
    cell.width = Cm(4)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    r1 = p.add_run(num + '\n')
    r1.bold = True
    r1.font.size = Pt(28)
    r1.font.name = 'Calibri'
    r1.font.color.rgb = rgb(99, 102, 241)
    r2 = p.add_run(lbl)
    r2.font.size = Pt(8.5)
    r2.font.name = 'Calibri'
    r2.font.color.rgb = rgb(148, 163, 184)
    p.paragraph_format.space_after = Pt(8)

doc.add_paragraph()
page_break(doc)

# ═══════════════════════════════════════ SECCIÓN 1 ══════════════════════════

heading(doc, '1. Estado Actual del Sistema', level=1)
divider(doc)
body(doc, 'El sistema implementa correctamente las operaciones core de un punto de venta moderno. '
     'A continuación el inventario completo de lo que ya existe y su nivel de completitud:')

rows_estado = [
    ['Venta en POS',              'Completo',    'Scan, búsqueda, descuentos, 3 listas de precio, botones rápidos'],
    ['Pagos combinados',          'Completo',    'Efectivo, débito, crédito, QR, transferencia, fiado, mixto'],
    ['Sesiones de caja (Multicaja)','Completo',  'Múltiples sesiones simultáneas, apertura/cierre, movimientos, saldo esperado'],
    ['Gestión de productos',      'Completo',    '3 listas de precio, costo, stock, categorías, proveedor, vencimiento'],
    ['Stock y ajustes',           'Completo',    'Ajuste manual, movimientos, alertas de stock bajo, reposición por proveedor'],
    ['Clientes / Cta. Corriente', 'Completo',    'Saldo, historial, pagos parciales, IVA para RI'],
    ['Proveedores',               'Completo',    'Órdenes de compra, recepción de mercadería, auto-órdenes'],
    ['Presupuestos',              'Completo',    'Con estados, vencimiento, impresión'],
    ['Devoluciones',              'Completo',    'Por venta, restock automático'],
    ['Rubros / Categorías',       'Completo',    'CRUD con conteo de productos por categoría'],
    ['Usuarios y roles',          'Completo',    'Admin, supervisor, cajero con permisos diferenciados'],
    ['Auditoría',                 'Completo',    'Log de acciones con usuario, entidad y detalle'],
    ['Backup',                    'Completo',    'Backup manual de la base SQLite'],
    ['Promociones',               'Básico',      'Porcentaje por producto — faltan condiciones temporales y por cantidad'],
    ['Reportes',                  'Básico',      'Resumen, ventas, márgenes, stock — sin comparativas ni tendencias'],
    ['Dashboard / BI',            'AUSENTE',     'No existe pantalla de resumen ni motor de insights'],
    ['Metas y proyecciones',      'AUSENTE',     'No hay sistema de objetivos ni proyección de cierre de mes'],
    ['CRM inteligente',           'AUSENTE',     'Clientes como base de datos, no como relaciones con análisis RFM'],
]
add_styled_table(doc, ['Módulo', 'Estado', 'Observaciones'], rows_estado, [4.5, 2.2, 9.8])

page_break(doc)

# ═══════════════════════════════════════ SECCIÓN 2 ══════════════════════════

heading(doc, '2. Diagnóstico: Puntos Débiles', level=1)
divider(doc)
body(doc, 'Los siguientes problemas limitan el valor real que el sistema entrega al dueño del negocio:')

colored_box(doc,
    '🔴  Reportes: números muertos sin contexto',
    'Los reportes muestran totales planos. No responden la pregunta real del dueño: "¿Estoy creciendo '
    'o cayendo?". Sin comparación período vs período, sin tendencia visual, sin benchmark interno. '
    'El dueño tiene que ir a 5 pestañas distintas para armar manualmente una foto del negocio.',
    bg='FEF2F2', title_color=rgb(153, 27, 27), text_color=rgb(127, 29, 29)
)
colored_box(doc,
    '🔴  El sistema no habla primero',
    'Si las ventas cayeron 30% esta semana, el sistema no dice nada. Si un producto se está por agotar, '
    'no avisa hasta que ya llegó al mínimo fijo. El dueño solo recibe información cuando la busca activamente. '
    'Todo el valor de los datos históricos está dormido.',
    bg='FEF2F2', title_color=rgb(153, 27, 27), text_color=rgb(127, 29, 29)
)
colored_box(doc,
    '🟡  Stock: reactivo, no predictivo',
    'La alerta de stock bajo llega cuando el producto ya llegó al mínimo — que el dueño ingresó a ojo '
    'y nunca actualizó. No se calcula velocidad de venta ni se estima cuántos días de stock quedan al '
    'ritmo actual de ventas.',
    bg='FFFBEB', title_color=rgb(146, 64, 14), text_color=rgb(120, 53, 15)
)
colored_box(doc,
    '🟡  Clientes: base de datos, no relaciones',
    'Solo existe nombre + saldo de deuda. No hay análisis de recencia ni frecuencia. No detecta que '
    'un cliente habitual dejó de venir. No hay segmentación. Los mejores clientes del negocio son '
    'completamente invisibles para el sistema.',
    bg='FFFBEB', title_color=rgb(146, 64, 14), text_color=rgb(120, 53, 15)
)
colored_box(doc,
    '🟡  Precios: sin inteligencia',
    'Los precios se actualizan a mano, producto por producto. No hay alerta de margen negativo. '
    'No se sabe qué 20% de productos genera el 80% de la ganancia. '
    'Actualizar por inflación requiere tocar cientos de productos individualmente.',
    bg='FFFBEB', title_color=rgb(146, 64, 14), text_color=rgb(120, 53, 15)
)
colored_box(doc,
    '🔵  Flujos operativos lentos',
    'Recepción de mercadería es ajuste manual de stock uno por uno. No hay importación masiva de '
    'productos por CSV. No hay conteo de inventario asistido. No hay cierre de día automático con informe.',
    bg='EFF6FF', title_color=rgb(30, 64, 175), text_color=rgb(30, 27, 75)
)

page_break(doc)

# ═══════════════════════════════════════ SECCIÓN 3 ══════════════════════════

heading(doc, '3. El Motor de Inteligencia de Negocio', level=1)
divider(doc)
body(doc, 'Todo lo siguiente es calculable sin ningún dato nuevo, únicamente con lo que ya está en la '
     'base de datos actual (ventas con timestamps, productos con costos, stock, clientes, movimientos):')

heading(doc, '3.1  Qué se puede calcular hoy mismo', level=2)

items_bi = [
    ('🕐 Distribución horaria de ventas', 'Cada venta tiene timestamp. Se mapea exactamente qué horas concentran mayor actividad y cuáles están muertas.'),
    ('📅 Patrón por día de semana', 'Promedio de ventas por día de semana + desviación del día actual respecto al histórico de ese mismo día.'),
    ('⚡ Velocidad de venta (unidades/día)', 'Unidades vendidas promedio por día en los últimos 30 días por producto. Permite calcular días de stock restantes.'),
    ('💀 Stock muerto', 'Productos con stock > 0 que no aparecen en ventas en los últimos N días. Dinero inmovilizado sin rotar.'),
    ('👥 RFM de clientes', 'Recencia (días desde última compra), Frecuencia (intervalo entre compras) y Monto (gasto promedio por visita).'),
    ('📈 Tendencia de margen', 'Ganancia bruta real por período, por categoría y por producto. Identificación de los SKUs que sostienen el negocio.'),
]
for icon_title, desc in items_bi:
    bullet(doc, f'{desc}', bold_prefix=f'{icon_title}: ')

heading(doc, '3.2  Los 23 insights automáticos posibles', level=2)
body(doc, 'Cada insight tiene: tipo, categoría, mensaje legible con números concretos, '
     'acción sugerida directa y confidence score basado en cantidad de datos históricos disponibles.')

rows_insights = [
    ['1',  'Producto X se agota en N días al ritmo actual',                         'Stock',       'URGENTE'],
    ['2',  'Stock muerto: $X inmovilizados en Y productos sin vender en 30+ días',  'Stock',       'IMPORTANTE'],
    ['3',  'Stock mínimo desactualizado en N productos (velocidad cambió)',          'Stock',       'IMPORTANTE'],
    ['4',  'Hoy es [día]. Históricamente tus [días] son X% mejores/peores',         'Ventas',      'CONSEJO'],
    ['5',  'Llevas $X hoy, estás Y% por debajo/encima de tu promedio del [día]',    'Ventas',      'INFO'],
    ['6',  'Este mes bajó/subió X% vs el mes pasado',                               'Ventas',      'IMPORTANTE'],
    ['7',  'Producto X aumentó N% esta semana — asegurate de tener stock',          'Ventas',      'CONSEJO'],
    ['8',  'Producto Y lleva 20+ días sin venderse — considerá una oferta',         'Ventas',      'CONSEJO'],
    ['9',  'Cliente X no compra hace N días (solía venir cada Y días)',             'Clientes',    'IMPORTANTE'],
    ['10', 'N clientes con CC no compraron en 30+ días y deben $X',                'Clientes',    'URGENTE'],
    ['11', 'La deuda del cliente X creció N% este mes',                            'Clientes',    'IMPORTANTE'],
    ['12', 'Tus 5 clientes más valiosos generan X% de tus ingresos',              'Clientes',    'INFO'],
    ['13', 'Producto X tiene margen negativo o inferior al mínimo configurado',    'Precios',     'URGENTE'],
    ['14', 'Categoría X tiene margen promedio del X%, muy debajo del resto',       'Precios',     'IMPORTANTE'],
    ['15', '20 productos generan el 80% de tu ganancia bruta (Pareto)',           'Precios',     'INFO'],
    ['16', 'Costos de N productos subieron vs última compra, precio no actualizado','Precios',    'IMPORTANTE'],
    ['17', 'Tu horario pico es H1-H2. Los [días] a esa hora vendés X% más',       'Operaciones', 'CONSEJO'],
    ['18', 'Hora muerta detectada: [H1-H2] tiene muy poca actividad',             'Operaciones', 'CONSEJO'],
    ['19', 'Pagos digitales: X% de tus ventas ya son por QR/transferencia',       'Operaciones', 'INFO'],
    ['20', 'Vas a necesitar recomprar X, Y, Z en los próximos 7 días',            'Compras',     'IMPORTANTE'],
    ['21', 'Proveedor X demora promedio N días en entregar',                       'Compras',     'INFO'],
    ['22', 'Si alcanzás $X hoy, cerrás el mes al N% de tu meta',                  'Metas',       'CONSEJO'],
    ['23', 'Producto con vencimiento próximo: X unidades vencen en N días',       'Stock',       'URGENTE'],
]
add_styled_table(doc, ['#', 'Insight generado automáticamente', 'Categoría', 'Urgencia'],
                 rows_insights, [0.7, 10.3, 2.5, 2.5])

page_break(doc)

# ═══════════════════════════════════════ SECCIÓN 4 ══════════════════════════

heading(doc, '4. Dashboard Ejecutivo', level=1)
divider(doc)
body(doc, 'Hoy cuando el dueño abre la app ve la pantalla de Caja. Debería ver primero un dashboard '
     'de 30 segundos que le cuente cómo está el negocio antes de empezar a vender.')

bloques = [
    ('Bloque 1 — Hoy', 'Ventas acumuladas vs mismo día semana pasada · cantidad de transacciones · ticket promedio · ganancia bruta estimada del día.'),
    ('Bloque 2 — Alertas urgentes', 'Centro del "inbox" del negocio: stock crítico (se agota en <2 días), productos próximos a vencer, cuentas corrientes en riesgo.'),
    ('Bloque 3 — Tendencia semanal', 'Mini gráfico de barras con ventas de los últimos 7 días + comparativa vs semana anterior (flecha ↑↓ + %).'),
    ('Bloque 4 — Top del día', 'Los 3 productos más vendidos hoy con unidades y monto. El método de pago dominante.'),
    ('Bloque 5 — Insights del día', '1 a 3 recomendaciones generadas automáticamente por el motor de análisis, priorizadas por urgencia.'),
    ('Bloque 6 — Progreso hacia meta', 'Barra de progreso hacia la meta diaria y mensual. Proyección de cierre de mes al ritmo actual.'),
]
for title, desc in bloques:
    bullet(doc, desc, bold_prefix=f'{title}: ')

# ═══════════════════════════════════════ SECCIÓN 5 ══════════════════════════

heading(doc, '5. El "Consejo del Día" — La Killer Feature', level=1)
divider(doc)
body(doc, 'La característica más diferenciadora del sistema: un briefing personalizado que se genera '
     'automáticamente cada mañana al abrir la app. Ejemplo real de cómo se vería:')

# Mockup box
mockup_tbl = doc.add_table(rows=1, cols=1)
mockup_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
cell = mockup_tbl.cell(0, 0)
set_cell_bg(cell, '0F172A')
cell.width = Inches(6.3)

lines = [
    ('☀️  Punto Simple — Buenos días, martes 27 de mayo', True, rgb(165, 180, 252), 11),
    ('─' * 72, False, rgb(30, 58, 95), 8),
    ('📊  Ayer: $47.200 en ventas  (+8% vs el lunes pasado)', True, rgb(134, 239, 172), 10.5),
    ('      → 3 productos generaron el 62% de esa cifra', False, rgb(100, 116, 139), 10),
    ('      → Mejor hora: 18-19h  (23% del total del día)', False, rgb(100, 116, 139), 10),
    ('─' * 72, False, rgb(30, 58, 95), 8),
    ('🔴  URGENTE: Coca Cola 2.25L se agota en ~1.5 días', True, rgb(248, 113, 113), 10.5),
    ('      → 12 unidades · velocidad: 8/día · pedí hoy', False, rgb(100, 116, 139), 10),
    ('🟡  Pedro García debe $23.000 hace 45 días sin compras', False, rgb(251, 191, 36), 10.5),
    ('─' * 72, False, rgb(30, 58, 95), 8),
    ('💡  Es martes. Tus martes promedian $38.000. Hoy llevas $0.', False, rgb(147, 197, 253), 10.5),
    ('💡  Galletitas Surtidas: 60 unidades sin vender en 3 semanas.', False, rgb(147, 197, 253), 10.5),
    ('      → 2x1 o 15% off podría recuperar $12.000 en efectivo', False, rgb(100, 116, 139), 10),
    ('🎯  Si llegás a $42.000 hoy → cerrás el mes al 94% de tu meta', True, rgb(134, 239, 172), 10.5),
]

first = True
for text, bold, color, size in lines:
    if first:
        p = cell.paragraphs[0]
        first = False
    else:
        p = cell.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Courier New'
    run.font.color.rgb = color

cell.paragraphs[-1].paragraph_format.space_after = Pt(8)
doc.add_paragraph()

body(doc, 'Este resumen se genera automáticamente al abrir la app, sin que el dueño tenga que ir '
     'a buscar nada. Toda la información ya existe en la base de datos actual.',
     italic=True)

page_break(doc)

# ═══════════════════════════════════════ SECCIÓN 6 ══════════════════════════

heading(doc, '6. Inteligencia de Stock', level=1)
divider(doc)

heading(doc, '6.1  De stock mínimo fijo a stock dinámico', level=2)
body(doc, 'El sistema actual usa un número fijo que el dueño ingresó una sola vez. El sistema inteligente calcula:')
bullet(doc, 'Unidades promedio vendidas por día en los últimos 30 días. Se actualiza con cada venta.',
       bold_prefix='Velocidad de venta: ')
bullet(doc, 'Stock actual ÷ velocidad diaria = días hasta agotamiento. Alerta configurable (ej: avisar con 3 días de anticipación).',
       bold_prefix='Días de stock restantes: ')
bullet(doc, 'Lead time del proveedor × velocidad de venta = punto de reorden real y dinámico, siempre actualizado.',
       bold_prefix='Stock mínimo sugerido: ')

heading(doc, '6.2  Detección de stock muerto', level=2)
colored_box(doc,
    '💀  Dinero inmovilizado sin rotar',
    'Productos con stock > 0 que no aparecen en ventas en los últimos 30/60/90 días. '
    'El sistema calcula cuánto capital está inmovilizado y sugiere acciones: descuento, '
    'liquidación, o devolución al proveedor.',
    bg='FFFBEB', title_color=rgb(146, 64, 14), text_color=rgb(120, 53, 15)
)

heading(doc, '6.3  Flujos operativos faltantes', level=2)
bullet(doc, 'Flujo guiado de conteo: el sistema presenta cada producto, el usuario ingresa la cantidad contada, '
       'y al final genera un informe de diferencias con ajuste masivo.', bold_prefix='Conteo de inventario asistido: ')
bullet(doc, 'En vez de ajustar stock uno por uno, un flujo de "Recibí X bultos de Proveedor Y" que '
       'actualiza stock en lote, vinculado a la orden de compra correspondiente.',
       bold_prefix='Recepción de mercadería rápida: ')

# ═══════════════════════════════════════ SECCIÓN 7 ══════════════════════════

heading(doc, '7. CRM: Inteligencia de Clientes', level=1)
divider(doc)

heading(doc, '7.1  Análisis RFM automático', level=2)
body(doc, 'Con los datos de ventas existentes el sistema puede clasificar a todos los clientes según:')
add_styled_table(doc,
    ['Dimensión', 'Qué mide', 'Cómo se calcula', 'Uso principal'],
    [
        ['Recencia (R)', 'Cuándo fue la última compra', 'Días desde la última venta con ese client_id', 'Detectar clientes ausentes'],
        ['Frecuencia (F)', 'Con qué regularidad compra', 'Promedio de días entre compras históricas', 'Detectar cambio de comportamiento'],
        ['Monto (M)', 'Cuánto gasta por visita', 'Total histórico ÷ cantidad de visitas', 'Segmentar por valor económico'],
    ],
    [3.5, 4, 5, 4]
)

heading(doc, '7.2  Segmentación automática de clientes', level=2)
segmentos = [
    ('⭐ VIP', 'Alto monto + alta frecuencia + recencia reciente. Los que más importan. Merecen atención especial y límites de crédito generosos.'),
    ('✅ Habitual', 'Frecuencia regular, gasto medio. El grueso del negocio. Retenerlos es la prioridad número uno.'),
    ('⚠️  En riesgo', 'Habituales que no vienen hace más tiempo del usual. Ventana de retención activa — el sistema alerta al dueño.'),
    ('🔴 Deudor crítico', 'Cuenta corriente alta + sin compras recientes. Doble señal de alerta financiera inmediata.'),
    ('👋 Ocasional', 'Baja frecuencia histórica. Candidatos para acciones de activación con promociones.'),
]
for name, desc in segmentos:
    bullet(doc, desc, bold_prefix=f'{name}: ')

page_break(doc)

# ═══════════════════════════════════════ SECCIÓN 8 ══════════════════════════

heading(doc, '8. Inteligencia de Precios y Márgenes', level=1)
divider(doc)

heading(doc, '8.1  Análisis Pareto 80/20 aplicado al negocio', level=2)
colored_box(doc,
    '📊  La regla 80/20 sobre los datos reales',
    'Con costo, precio y ventas históricas se calcula exactamente qué productos generan el 80% '
    'de la ganancia bruta. En la mayoría de los negocios esto es sorprendente: 15-20 SKUs sostienen '
    'el negocio. Los otros 200 productos hacen ruido. Conocer esto permite tomar decisiones '
    'de surtido, espacio y atención completamente distintas.',
    bg='EFF6FF', title_color=rgb(30, 64, 175), text_color=rgb(30, 27, 75)
)

heading(doc, '8.2  Actualización masiva de precios por inflación', level=2)
bullet(doc, 'Aplicar un % de aumento a toda la tienda, una categoría, o un proveedor específico.', bold_prefix='Actualización en lote: ')
bullet(doc, 'Aplicar distintos % según el margen actual: los de menor margen suben más, los de mayor margen se tocan menos.', bold_prefix='Aumento diferenciado: ')
bullet(doc, 'Ver exactamente qué precios cambian y cuánto antes de confirmar, con opción de exclusiones.', bold_prefix='Preview antes de confirmar: ')
bullet(doc, 'El sistema recuerda cuándo fue la última actualización por categoría.', bold_prefix='Registro de actualizaciones: ')

heading(doc, '8.3  Otras features de precios', level=2)
bullet(doc, 'En la pantalla de venta, si el cajero aplica un descuento grande, alertar si el margen cae por debajo del mínimo configurado.', bold_prefix='Alerta de margen en POS: ')
bullet(doc, 'Cuando llega una orden de compra con precios más altos, detectar automáticamente los productos donde el costo subió pero el precio no.', bold_prefix='Detección de desajuste costo-precio: ')
bullet(doc, 'Configurar el costo de cada método de pago (ej: 2.5% QR, 3% tarjeta) para calcular la rentabilidad real de cada venta.', bold_prefix='Comisiones por método de pago: ')

# ═══════════════════════════════════════ SECCIÓN 9 ══════════════════════════

heading(doc, '9. Inteligencia de Compras y Proveedores', level=1)
divider(doc)

bullet(doc, 'Las órdenes ya tienen fecha de creación. Agregar fecha de recepción permite calcular el tiempo real de entrega. Con ese dato, el punto de reorden se vuelve preciso.', bold_prefix='Lead time real por proveedor: ')
bullet(doc, '"Basado en tu velocidad de venta, vas a necesitar recomprar estos productos en los próximos 7 días. Costo estimado: $87.000." Con botón para generar la orden directamente.', bold_prefix='Proyección de compras a 7 días: ')
bullet(doc, 'Comparando el precio en cada orden de compra, mostrar la inflación real de cada proveedor mes a mes.', bold_prefix='Detección de inflación de costos: ')
bullet(doc, 'Si el mismo producto puede venir de dos proveedores, mostrar el historial de precios de cada uno al momento de hacer el pedido.', bold_prefix='Comparativa entre proveedores: ')

# ═══════════════════════════════════════ SECCIÓN 10 ═════════════════════════

heading(doc, '10. Metas, Rentabilidad Real y Punto de Equilibrio', level=1)
divider(doc)

heading(doc, '10.1  Sistema de metas', level=2)
bullet(doc, 'Meta de ventas diaria (ej: $50.000/día)')
bullet(doc, 'Meta mensual (ej: $1.200.000/mes)')
bullet(doc, 'Margen mínimo objetivo por categoría')
body(doc, 'El dashboard muestra progreso en tiempo real con barra visual y proyección de cierre '
     'de mes al ritmo actual.')

heading(doc, '10.2  Costos operativos y rentabilidad neta', level=2)
colored_box(doc,
    '💰  El dato más importante que hoy falta',
    'El sistema conoce el costo de los productos pero no los costos operativos (alquiler, salarios, servicios). '
    'Sin ese dato, la "ganancia" que muestra es bruta. Agregar una sección de gastos fijos mensuales '
    'permite calcular la ganancia neta real del negocio.',
    bg='FFFBEB', title_color=rgb(146, 64, 14), text_color=rgb(120, 53, 15)
)

heading(doc, '10.3  Punto de equilibrio diario', level=2)
bullet(doc, 'Punto de equilibrio diario = costos fijos mensuales ÷ días hábiles')
bullet(doc, 'El dashboard muestra: "Necesitás vender $28.000 hoy para cubrir costos fijos"')
bullet(doc, 'En tiempo real: "Cubriste el punto de equilibrio a las 14:32h ✓"')

page_break(doc)

# ═══════════════════════════════════════ SECCIÓN 11 ════════════════════════

heading(doc, '11. Mejoras de Usabilidad y Flujos Faltantes', level=1)
divider(doc)

items_ux = [
    ('Importación masiva de productos (CSV)', 'Subir miles de productos de una vez con mapeador de columnas visual, validación previa y deduplicación por código de barras. Esencial para negocios que cambian proveedor o actualizan catálogos completos.'),
    ('Búsqueda difusa (fuzzy search)', 'Si el usuario escribe "coc" debería encontrar "Coca Cola". Si escribe "coac" también. Tolerancia a errores tipográficos mejora drásticamente la velocidad en el POS.'),
    ('Cierre de día automático con PDF', 'Al cerrar la última sesión de caja, generar automáticamente un "Resumen del día" en PDF: ventas por método de pago, top 5 productos, comparativa, ganancia bruta, movimientos de caja.'),
    ('Impresión de etiquetas de precio', 'Desde la pantalla de productos, seleccionar uno o varios y generar etiquetas (código de barras + nombre + precio) para impresora de etiquetas o hoja A4.'),
    ('Combos y packs', 'Vender X + Y juntos a un precio especial. El sistema descuenta el stock de cada componente individualmente pero cobra el precio del combo.'),
    ('Promociones con condiciones temporales', 'Promociones activas solo en ciertos días u horarios, "2da unidad al X%", "a partir de 3 unidades X% de descuento", "llevá A + B por $X".'),
    ('Modo consulta rápida tablet/móvil', 'Vista simplificada en la misma red local, accesible desde el teléfono del dueño para ver ventas del día, alertas urgentes y aprobar órdenes sin estar en la PC.'),
    ('Libro IVA para Responsable Inscripto', 'Listado de ventas con IVA discriminado (neto + 21%) agrupado por período fiscal, listo para el contador.'),
]
for title, desc in items_ux:
    bullet(doc, desc, bold_prefix=f'{title}: ')

# ═══════════════════════════════════════ SECCIÓN 12 ════════════════════════

heading(doc, '12. Exportaciones y Contabilidad', level=1)
divider(doc)

body(doc, 'Hoy no hay forma de sacar los datos en formatos útiles para el contador o para análisis externo:')

add_styled_table(doc,
    ['Exportación', 'Contenido', 'Destinatario'],
    [
        ['Libro IVA Ventas', 'Ventas con IVA discriminado (neto + 21%), agrupado por período fiscal', 'Contador RI'],
        ['Ventas en CSV/Excel', 'Todas las ventas del período: fecha, total, método de pago, cliente, productos', 'Contador / análisis'],
        ['Movimientos de stock', 'Entradas, salidas, ajustes con fechas y motivos', 'Contador / auditoría'],
        ['Cuentas corrientes', 'Saldo y movimientos de todos los clientes', 'Cobros / legal'],
        ['Informe mensual PDF', 'Ventas, costo de ventas, ganancia bruta, desglose por categoría, cuentas por cobrar', 'Contador / dueño'],
        ['Conciliación de caja', 'Movimientos del período separados por efectivo, transferencias y QR', 'Conciliación bancaria'],
    ],
    [4.5, 8.5, 3.5]
)

page_break(doc)

# ═══════════════════════════════════════ SECCIÓN 13 ════════════════════════

heading(doc, '13. Hoja de Ruta — Roadmap Priorizado', level=1)
divider(doc)

body(doc, 'Cuatro fases ordenadas por impacto/esfuerzo. Las primeras dos fases se ejecutan sobre datos ya existentes.')

add_styled_table(doc,
    ['Fase', 'Feature', 'Impacto', 'Esfuerzo'],
    [
        ['FASE 1 — Quick Wins', 'Dashboard con métricas del día + comparativa período vs período', 'MUY ALTO', 'BAJO'],
        ['FASE 1', 'Centro de alertas proactivas (stock, deuda, vencimientos)', 'MUY ALTO', 'BAJO'],
        ['FASE 1', 'Velocidad de venta + días de stock restantes', 'ALTO', 'BAJO'],
        ['FASE 1', 'Alerta de margen negativo / insuficiente por producto', 'ALTO', 'BAJO'],
        ['FASE 1', 'Exportación CSV / informe mensual PDF', 'ALTO', 'BAJO'],
        ['FASE 2 — Inteligencia', 'Motor de 23 insights automáticos', 'MUY ALTO', 'MEDIO'],
        ['FASE 2', 'Briefing diario / "Consejo del Día"', 'MUY ALTO', 'MEDIO'],
        ['FASE 2', 'Análisis RFM + segmentación automática de clientes', 'ALTO', 'MEDIO'],
        ['FASE 2', 'Análisis Pareto 80/20 de rentabilidad por producto', 'ALTO', 'BAJO'],
        ['FASE 2', 'Detección costo sube / precio no actualizado', 'ALTO', 'BAJO'],
        ['FASE 3 — Profundidad', 'Sistema de metas + punto de equilibrio + rentabilidad neta', 'ALTO', 'MEDIO'],
        ['FASE 3', 'Actualización masiva de precios por inflación', 'ALTO', 'MEDIO'],
        ['FASE 3', 'Importación masiva de productos por CSV', 'ALTO', 'MEDIO'],
        ['FASE 3', 'Proyección de compras a 7 días + lead time real', 'ALTO', 'MEDIO'],
        ['FASE 3', 'Conteo de inventario asistido + recepción rápida de mercadería', 'MEDIO', 'MEDIO'],
        ['FASE 3', 'Cierre de día automático con PDF', 'MEDIO', 'BAJO'],
        ['FASE 4 — Expansión', 'Combos y packs con stock por componente', 'MEDIO', 'ALTO'],
        ['FASE 4', 'Promociones temporales (horario, día, cantidad)', 'MEDIO', 'ALTO'],
        ['FASE 4', 'Modo consulta rápida tablet/red local', 'MEDIO', 'ALTO'],
        ['FASE 4', 'Facturación electrónica AFIP', 'ALTO', 'MUY ALTO'],
    ],
    [3.0, 9.0, 2.5, 2.0]
)

# ═══════════════════════════════════════ SECCIÓN 14 ════════════════════════

heading(doc, '14. Catálogo Completo de Features (42 en total)', level=1)
divider(doc)

all_features = [
    ('1', 'Dashboard ejecutivo con métricas del día', 'Dashboard', '1'),
    ('2', 'Comparativa período vs período en reportes', 'Reportes', '1'),
    ('3', 'Centro de alertas proactivas', 'Dashboard', '1'),
    ('4', 'Velocidad de venta + días de stock restantes', 'Stock', '1'),
    ('5', 'Exportación CSV / PDF informe mensual', 'Reportes', '1'),
    ('6', 'Alerta de margen negativo / insuficiente', 'Precios', '1'),
    ('7', 'Detección costo subió / precio no actualizado', 'Precios', '1'),
    ('8', 'Motor de 23 insights automáticos', 'BI Engine', '2'),
    ('9', 'Briefing diario / "Consejo del Día"', 'BI Engine', '2'),
    ('10', 'Análisis RFM de clientes', 'Clientes', '2'),
    ('11', 'Segmentación automática (VIP / habitual / riesgo)', 'Clientes', '2'),
    ('12', 'Detección de clientes habituales ausentes', 'Clientes', '2'),
    ('13', 'Análisis Pareto 80/20 de rentabilidad', 'Reportes', '2'),
    ('14', 'Detección de stock muerto con costo inmovilizado', 'Stock', '2'),
    ('15', 'Distribución horaria de ventas', 'Reportes', '2'),
    ('16', 'Sistema de metas diarias y mensuales', 'Dashboard', '3'),
    ('17', 'Costos operativos y rentabilidad neta', 'Finanzas', '3'),
    ('18', 'Punto de equilibrio diario', 'Finanzas', '3'),
    ('19', 'Proyección de cierre de mes (3 escenarios)', 'Finanzas', '3'),
    ('20', 'Actualización masiva de precios por inflación', 'Productos', '3'),
    ('21', 'Importación masiva de productos por CSV', 'Productos', '3'),
    ('22', 'Proyección de compras a 7 días', 'Proveedores', '3'),
    ('23', 'Lead time real por proveedor', 'Proveedores', '3'),
    ('24', 'Flujo de recepción de mercadería rápida', 'Stock', '3'),
    ('25', 'Conteo de inventario asistido', 'Stock', '3'),
    ('26', 'Cierre de día automático con PDF', 'Caja', '3'),
    ('27', 'Libro IVA Ventas para Responsable Inscripto', 'Contabilidad', '3'),
    ('28', 'Búsqueda difusa con tolerancia a errores', 'POS', '3'),
    ('29', 'Alerta de margen al aplicar descuento en POS', 'POS', '3'),
    ('30', 'Detección de inflación de costos por proveedor', 'Proveedores', '3'),
    ('31', 'Comparativa de precios entre proveedores', 'Proveedores', '3'),
    ('32', 'Combos y packs con stock por componente', 'POS', '4'),
    ('33', 'Promociones temporales (horario, día, cantidad)', 'Promociones', '4'),
    ('34', 'Impresión de etiquetas de precio', 'Productos', '4'),
    ('35', 'Modo consulta rápida tablet/móvil', 'UX', '4'),
    ('36', 'Backup automático programado', 'Sistema', '4'),
    ('37', 'Comisiones por método de pago', 'Finanzas', '4'),
    ('38', 'Notas y recordatorios en clientes', 'Clientes', '4'),
    ('39', 'Límites de crédito automáticos sugeridos', 'Clientes', '4'),
    ('40', 'Historial de precios de costo por producto', 'Productos', '4'),
    ('41', 'Facturación electrónica AFIP', 'Fiscal', 'Futuro'),
    ('42', 'Multi-dispositivo en red local', 'Sistema', 'Futuro'),
]
add_styled_table(doc,
    ['#', 'Feature', 'Módulo', 'Fase'],
    all_features,
    [0.7, 10.8, 3.0, 2.0]
)

page_break(doc)

# ═══════════════════════════════════════ SECCIÓN 15 ════════════════════════

heading(doc, '15. Matriz de Impacto vs Esfuerzo', level=1)
divider(doc)
body(doc, 'Cada feature clasificada en cuatro cuadrantes para priorización estratégica:')

matrix_tbl = doc.add_table(rows=2, cols=2)
matrix_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

quadrants = [
    ('ALTO IMPACTO · BAJO ESFUERZO\nHACER PRIMERO ✓', 'DCFCE7', rgb(22, 101, 52), [
        'Dashboard con métricas del día',
        'Comparativa período vs período',
        'Velocidad de venta + días de stock',
        'Alertas proactivas (stock, deuda)',
        'Detección de margen negativo',
        'Exportación CSV / informe PDF',
        'Detección costo sube / precio no',
    ]),
    ('ALTO IMPACTO · ESFUERZO MEDIO\nPLANIFICAR', 'DBEAFE', rgb(29, 78, 216), [
        'Motor de 23 insights automáticos',
        '"Consejo del Día" / briefing matutino',
        'Análisis RFM + segmentación clientes',
        'Sistema de metas y punto equilibrio',
        'Actualización masiva de precios',
        'Proyección de compras a 7 días',
        'Análisis Pareto 80/20',
    ]),
    ('IMPACTO MEDIO · BAJO ESFUERZO\nLLENAR HUECOS', 'EDE9FE', rgb(109, 40, 217), [
        'Cierre de día automático PDF',
        'Alerta de margen al descontar en POS',
        'Libro IVA para RI',
        'Impresión de etiquetas',
        'Notas en clientes',
        'Lead time real por proveedor',
    ]),
    ('IMPACTO MEDIO · ESFUERZO ALTO\nEVALUAR', 'FEF9C3', rgb(133, 77, 14), [
        'Importación masiva CSV',
        'Combos y packs',
        'Modo tablet / red local',
        'Facturación electrónica AFIP',
        'Backup automático en nube',
        'Búsqueda difusa',
    ]),
]

positions = [(0,0), (0,1), (1,0), (1,1)]
for (r, c), (title, bg, title_col, items_list) in zip(positions, quadrants):
    cell = matrix_tbl.cell(r, c)
    set_cell_bg(cell, bg)
    cell.width = Cm(8.0)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Calibri'
    run.font.color.rgb = title_col
    for item in items_list:
        pi = cell.add_paragraph()
        pi.paragraph_format.left_indent = Cm(0.5)
        pi.paragraph_format.space_before = Pt(1)
        pi.paragraph_format.space_after = Pt(1)
        ri = pi.add_run(f'• {item}')
        ri.font.size = Pt(9.5)
        ri.font.name = 'Calibri'
        ri.font.color.rgb = rgb(71, 85, 105)
    cell.paragraphs[-1].paragraph_format.space_after = Pt(8)

doc.add_paragraph()

# Conclusión final
colored_box(doc,
    '🎯  Conclusión Estratégica',
    'El sistema tiene una base técnica sólida. La inversión para el próximo ciclo debe concentrarse en '
    'las 7 features de alto impacto / bajo esfuerzo: todas son calculables con datos ya existentes y '
    'transforman inmediatamente la propuesta de valor del producto.\n\n'
    'El "Consejo del Día" y el motor de insights son la diferenciación central frente a cualquier '
    'otro POS del mercado: ningún competidor convierte los datos del negocio en recomendaciones '
    'accionables y personalizadas. Esa es la ventaja competitiva real.',
    bg='0F172A', title_color=rgb(165, 180, 252), text_color=rgb(203, 213, 225)
)

# ─── GUARDAR ────────────────────────────────────────────────────────────────

out_path = r'c:\Users\magal\OneDrive\Escritorio\kiosco-pos\kiosco-pos\DOCUMENTACIÓN\Analisis_Estrategico_PuntoSimple.docx'
doc.save(out_path)
print(f'Guardado: {out_path}')
